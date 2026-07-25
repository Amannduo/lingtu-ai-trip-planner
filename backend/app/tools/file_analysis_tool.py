"""File analysis agent — parse travel documents and analyze with LLM.

Supports: TXT, MD, PDF, DOCX, XLSX.
"""

from __future__ import annotations

import os
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any


MAX_EXTRACTED_CHARS = 200_000
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_ARCHIVE_FILES = 5_000


def _read_text(path: str) -> str:
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(path, "r", encoding=encoding) as fh:
                return fh.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    return ""


def _validate_office_archive(path: str) -> None:
    """Reject oversized or highly expanded Office ZIP containers."""
    if not zipfile.is_zipfile(path):
        raise ValueError("Office 文件不是有效的 ZIP 容器")
    with zipfile.ZipFile(path) as archive:
        members = archive.infolist()
        if len(members) > MAX_ARCHIVE_FILES:
            raise ValueError("Office 文件包含过多条目")
        total_size = sum(member.file_size for member in members)
        if total_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
            raise ValueError("Office 文件解压后过大")
        for member in members:
            compressed = max(1, member.compress_size)
            if member.file_size > 5 * 1024 * 1024 and member.file_size / compressed > 200:
                raise ValueError("Office 文件压缩比异常")


# ── File parsers ────────────────────────────────────────────────────────

def parse_txt(path: str) -> str:
    return _read_text(path)


def parse_md(path: str) -> str:
    return _read_text(path)


def parse_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[PDF 解析需要安装 pypdf]"

    reader = PdfReader(path)
    if len(reader.pages) > 200:
        raise ValueError("PDF 页数不能超过200页")
    pages: list[str] = []
    extracted = 0
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
            extracted += len(text)
            if extracted >= MAX_EXTRACTED_CHARS:
                break
    return "\n\n".join(pages)[:MAX_EXTRACTED_CHARS]


def parse_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[DOCX 解析需要安装 python-docx]"

    _validate_office_archive(path)
    doc = Document(path)
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # Also extract tables
    tables_text: list[str] = []
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables_text.append("\n".join(rows))

    result = "\n\n".join(paragraphs)
    if tables_text:
        result += "\n\n--- 表格数据 ---\n\n" + "\n\n".join(tables_text)
    return result[:MAX_EXTRACTED_CHARS]


def parse_xlsx(path: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[XLSX 解析需要安装 openpyxl]"

    _validate_office_archive(path)
    wb = load_workbook(path, data_only=True, read_only=True)
    output: list[str] = []
    extracted = 0
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        output.append(f"## 工作表: {sheet_name}")
        for row_index, row in enumerate(ws.iter_rows(values_only=True)):
            if row_index >= 10_000 or extracted >= MAX_EXTRACTED_CHARS:
                break
            cells = [str(cell)[:2_000] if cell is not None else "" for cell in row[:200]]
            non_empty = [cell for cell in cells if cell]
            if non_empty:
                line = " | ".join(cells)
                output.append(line)
                extracted += len(line)
        output.append("")
        if extracted >= MAX_EXTRACTED_CHARS:
            break
    wb.close()
    return "\n".join(output)[:MAX_EXTRACTED_CHARS]


# ── Extension → parser mapping ──────────────────────────────────────────

PARSERS = {
    ".txt": parse_txt,
    ".md": parse_md,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
}


def parse_uploaded_file(file_path: str) -> tuple[str, str]:
    """Parse an uploaded file. Returns (content, file_type_label)."""
    suffix = Path(file_path).suffix.lower()
    parser = PARSERS.get(suffix)
    if parser is None:
        return f"[不支持的文件类型: {suffix}]", f"未知类型 ({suffix})"
    content = parser(file_path)
    return content[:MAX_EXTRACTED_CHARS], suffix.lstrip(".").upper()


# ── LLM analysis ─────────────────────────────────────────────────────────

FILE_ANALYSIS_SYSTEM_PROMPT = """你是一个旅行文件分析助手。用户上传了文件，请分析其中的旅行相关内容。

请严格按以下 JSON 格式返回，不要输出其他内容：
{
  "summary": "文件内容摘要（2-5句中文）",
  "suggestions": ["建议1", "建议2", "建议3"],
  "extracted_info": {
    "cities": ["城市名"],
    "dates": ["日期"],
    "budget": "预算信息或空",
    "travelers": "人数或空"
  },
  "table": []
}

要求：
1. summary 概括文件中的旅行计划、行程要点
2. suggestions 给出 2-4 条改进建议（预算、时间、景点、交通等方面）
3. extracted_info 提取结构化数据
4. table 如果文件包含表格数据，转换为 JSON 数组格式
5. 只返回 JSON，不要输出 Markdown
"""


def analyze_travel_file(content: str, question: str | None = None) -> dict[str, Any]:
    """Send extracted file content to LLM for analysis."""
    from ..services.llm_service import get_llm
    from hello_agents import SimpleAgent

    llm = get_llm()
    agent = SimpleAgent(
        name="FileAnalyzer",
        llm=llm,
        system_prompt=FILE_ANALYSIS_SYSTEM_PROMPT,
    )

    user_prompt = f"请分析以下文件内容：\n\n{content[:8000]}"  # truncate for safety
    if question:
        user_prompt += f"\n\n用户额外提问：{question}"

    response = agent.run(user_prompt)
    return _parse_analysis_response(response, content)


def _parse_analysis_response(response: str, original_content: str) -> dict[str, Any]:
    """Parse LLM JSON response, with fallback."""
    import json as _json

    text = response.strip()
    # strip fences
    for marker in ("```json", "```"):
        if marker in text:
            start = text.find(marker) + len(marker)
            end = text.rfind("```")
            if end > start:
                text = text[start:end].strip()
            else:
                text = text[start:].strip()
            break
    brace_start = text.find("{")
    brace_end = text.rfind("}")
    if (
        not text.lstrip().startswith("[")
        and brace_start >= 0
        and brace_end > brace_start
    ):
        text = text[brace_start:brace_end + 1]

    try:
        parsed = _json.loads(text)
        if not isinstance(parsed, dict):
            raise ValueError("analysis response must be a JSON object")
        parsed["_analysis_degraded"] = False
        return parsed
    except (TypeError, ValueError, _json.JSONDecodeError):
        return {
            "_analysis_degraded": True,
            "summary": f"文件解析完成，共 {len(original_content)} 个字符。LLM 返回格式异常，以下是原始内容摘要。",
            "suggestions": ["请检查文件格式是否正确", "尝试使用更简洁的内容重新上传"],
            "extracted_info": {"cities": [], "dates": [], "budget": "", "travelers": ""},
            "table": [],
        }


# ── Convenience: parse + analyze ─────────────────────────────────────────

def process_uploaded_file(file_path: str, question: str | None = None) -> dict[str, Any]:
    """Parse a file and analyze it with LLM.  One-stop entry point."""
    content, file_type = parse_uploaded_file(file_path)

    if content.startswith("[不支持的") or content.startswith("[PDF 解析需要") or content.startswith("[DOCX 解析需要") or content.startswith("[XLSX 解析需要"):
        return {
            "success": False,
            "summary": content,
            "suggestions": [],
            "extracted_info": {},
            "table": [],
            "file_type": file_type,
        }

    analysis = analyze_travel_file(content, question)
    analysis["success"] = not bool(analysis.pop("_analysis_degraded", False))
    analysis["file_type"] = file_type
    return analysis
